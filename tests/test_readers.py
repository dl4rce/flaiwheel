# Flaiwheel – Self-improving knowledge base for AI coding agents
# Copyright (c) 2026 4rce.com Digital Technologies GmbH.
# Use of this software is governed by the Business Source License 1.1. See LICENSE.md.

"""Tests for multi-format file readers."""

import json
import textwrap
from pathlib import Path

import pytest

from flaiwheel.readers import (
    SUPPORTED_EXTENSIONS,
    extract_text,
    _read_md,
    _read_txt,
    _read_pdf,
    _read_html,
    _read_rst,
    _read_docx,
    _read_json,
    _read_yaml,
    _read_csv,
)


@pytest.fixture
def tmp(tmp_path):
    return tmp_path


class TestSupportedExtensions:
    def test_all_nine_formats(self):
        expected = {".md", ".txt", ".pdf", ".html", ".htm", ".rst",
                    ".docx", ".json", ".yaml", ".yml", ".csv"}
        assert SUPPORTED_EXTENSIONS == expected

    def test_unsupported_returns_none(self, tmp):
        f = tmp / "data.xlsx"
        f.write_text("some data")
        assert extract_text(f) is None

    def test_nonexistent_file_returns_none(self, tmp):
        f = tmp / "ghost.md"
        assert extract_text(f) is None


class TestMarkdown:
    def test_passthrough(self, tmp):
        f = tmp / "readme.md"
        f.write_text("# Hello\n\nWorld")
        result = extract_text(f)
        assert result == "# Hello\n\nWorld"

    def test_preserves_content(self, tmp):
        content = "# Title\n\n## Section\n\nSome text with `code`."
        f = tmp / "doc.md"
        f.write_text(content)
        assert extract_text(f) == content


class TestPlainText:
    def test_wraps_in_heading(self, tmp):
        f = tmp / "notes.txt"
        f.write_text("These are my notes.\nLine two.")
        result = extract_text(f)
        assert result.startswith("# notes.txt")
        assert "These are my notes." in result
        assert "Line two." in result

    def test_empty_file(self, tmp):
        f = tmp / "empty.txt"
        f.write_text("")
        result = extract_text(f)
        assert result is not None
        assert "# empty.txt" in result


class TestPDF:
    def test_extract_text(self, tmp):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)

        page = writer.pages[0]
        from io import BytesIO
        from pypdf.generic import (
            ArrayObject,
            DecodedStreamObject,
            DictionaryObject,
            NameObject,
            NumberObject,
        )
        font_dict = DictionaryObject()
        font_dict[NameObject("/Type")] = NameObject("/Font")
        font_dict[NameObject("/Subtype")] = NameObject("/Type1")
        font_dict[NameObject("/BaseFont")] = NameObject("/Helvetica")

        resources = page.get("/Resources", DictionaryObject())
        if "/Font" not in resources:
            resources[NameObject("/Font")] = DictionaryObject()
        resources["/Font"][NameObject("/F1")] = font_dict

        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 50 150 Td (Hello PDF World) Tj ET")
        page[NameObject("/Contents")] = stream

        pdf_path = tmp / "test.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        result = extract_text(pdf_path)
        assert result is not None
        assert "test.pdf" in result
        assert "Hello PDF World" in result

    def test_empty_pdf(self, tmp):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        pdf_path = tmp / "blank.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)

        result = extract_text(pdf_path)
        assert result is None


class TestHTML:
    def test_headings_converted(self, tmp):
        html = "<html><body><h1>Title</h1><h2>Sub</h2><p>Text</p></body></html>"
        f = tmp / "doc.html"
        f.write_text(html)
        result = extract_text(f)
        assert "# Title" in result
        assert "## Sub" in result
        assert "Text" in result

    def test_lists_converted(self, tmp):
        html = "<html><body><ul><li>Item A</li><li>Item B</li></ul></body></html>"
        f = tmp / "list.html"
        f.write_text(html)
        result = extract_text(f)
        assert "- Item A" in result
        assert "- Item B" in result

    def test_code_blocks(self, tmp):
        html = "<html><body><pre>print('hello')</pre></body></html>"
        f = tmp / "code.html"
        f.write_text(html)
        result = extract_text(f)
        assert "```" in result
        assert "print('hello')" in result

    def test_scripts_stripped(self, tmp):
        html = "<html><body><script>alert('x')</script><p>Real content</p></body></html>"
        f = tmp / "scripts.html"
        f.write_text(html)
        result = extract_text(f)
        assert "alert" not in result
        assert "Real content" in result

    def test_htm_extension(self, tmp):
        html = "<html><body><p>HTM file</p></body></html>"
        f = tmp / "doc.htm"
        f.write_text(html)
        result = extract_text(f)
        assert "HTM file" in result

    def test_empty_html(self, tmp):
        f = tmp / "empty.html"
        f.write_text("<html><body></body></html>")
        result = extract_text(f)
        assert result is None


class TestRST:
    def test_heading_underline(self, tmp):
        rst = textwrap.dedent("""\
            My Title
            ========

            Some paragraph text.

            Subsection
            ----------

            More text.
        """)
        f = tmp / "doc.rst"
        f.write_text(rst)
        result = extract_text(f)
        assert "# My Title" in result
        assert "## Subsection" in result
        assert "Some paragraph text." in result

    def test_code_block(self, tmp):
        rst = textwrap.dedent("""\
            Example
            =======

            .. code-block:: python

               print("hello")
               x = 1
        """)
        f = tmp / "code.rst"
        f.write_text(rst)
        result = extract_text(f)
        assert "```" in result
        assert 'print("hello")' in result

    def test_overline_title(self, tmp):
        rst = textwrap.dedent("""\
            =========
            My Title
            =========

            Content here.
        """)
        f = tmp / "overline.rst"
        f.write_text(rst)
        result = extract_text(f)
        assert "# My Title" in result


class TestDOCX:
    def test_paragraphs_and_headings(self, tmp):
        from docx import Document

        doc = Document()
        doc.add_heading("Document Title", level=1)
        doc.add_heading("Section One", level=2)
        doc.add_paragraph("This is paragraph text.")
        doc.add_heading("Section Two", level=3)
        doc.add_paragraph("More content here.")

        docx_path = tmp / "test.docx"
        doc.save(docx_path)

        result = extract_text(docx_path)
        assert result is not None
        assert "# Document Title" in result
        assert "## Section One" in result
        assert "### Section Two" in result
        assert "This is paragraph text." in result
        assert "More content here." in result

    def test_empty_docx(self, tmp):
        from docx import Document

        doc = Document()
        docx_path = tmp / "empty.docx"
        doc.save(docx_path)

        result = extract_text(docx_path)
        assert result is None


class TestJSON:
    def test_wraps_in_code_block(self, tmp):
        data = {"name": "test", "version": "1.0"}
        f = tmp / "config.json"
        f.write_text(json.dumps(data))
        result = extract_text(f)
        assert "# config.json" in result
        assert "```json" in result
        assert '"name": "test"' in result

    def test_pretty_prints(self, tmp):
        f = tmp / "compact.json"
        f.write_text('{"a":1,"b":2}')
        result = extract_text(f)
        assert '"a": 1' in result

    def test_invalid_json_passthrough(self, tmp):
        f = tmp / "broken.json"
        f.write_text("{not valid json}")
        result = extract_text(f)
        assert "```json" in result
        assert "{not valid json}" in result


class TestYAML:
    def test_wraps_in_code_block(self, tmp):
        f = tmp / "config.yaml"
        f.write_text("name: test\nversion: 1.0\n")
        result = extract_text(f)
        assert "# config.yaml" in result
        assert "```yaml" in result
        assert "name: test" in result

    def test_yml_extension(self, tmp):
        f = tmp / "data.yml"
        f.write_text("key: value\n")
        result = extract_text(f)
        assert "# data.yml" in result
        assert "```yaml" in result


class TestCSV:
    def test_markdown_table(self, tmp):
        f = tmp / "data.csv"
        f.write_text("Name,Age,City\nAlice,30,Berlin\nBob,25,Munich\n")
        result = extract_text(f)
        assert "# data.csv" in result
        assert "| Name | Age | City |" in result
        assert "| --- | --- | --- |" in result
        assert "| Alice | 30 | Berlin |" in result
        assert "| Bob | 25 | Munich |" in result

    def test_empty_csv(self, tmp):
        f = tmp / "empty.csv"
        f.write_text("")
        result = extract_text(f)
        assert "(empty)" in result

    def test_uneven_rows(self, tmp):
        f = tmp / "uneven.csv"
        f.write_text("A,B,C\n1,2\n")
        result = extract_text(f)
        assert "| 1 | 2 |" in result


class TestExtractTextIntegration:
    def test_all_formats_produce_output(self, tmp):
        """Verify that every format with simple content returns something."""
        files = {
            "test.md": "# Markdown",
            "test.txt": "Plain text",
            "test.html": "<h1>HTML</h1>",
            "test.rst": "RST Title\n=========\n\nText.",
            "test.json": '{"key": "value"}',
            "test.yaml": "key: value\n",
            "test.csv": "a,b\n1,2\n",
        }
        for name, content in files.items():
            f = tmp / name
            f.write_text(content)
            result = extract_text(f)
            assert result is not None, f"extract_text returned None for {name}"
            assert len(result) > 0, f"extract_text returned empty for {name}"
